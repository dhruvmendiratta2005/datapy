import cv2
import numpy as np
import pytesseract
from PIL import Image
from docx import Document

def do(filename, output_filename):
    # Load the image
    image = cv2.imread(filename)

    # Check if image was loaded
    if image is None:
        print("Error: Image not found.")
        return

    # Convert to grayscale
    gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    # Denoise the image
    denoised_image = cv2.fastNlMeansDenoising(gray_image)

    # Thresholding
    _, binary_image = cv2.threshold(denoised_image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    # Deskewing (optional)
    coords = np.column_stack(np.where(binary_image > 0))
    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle
    (h, w) = binary_image.shape[:2]
    center = (w // 2, h // 2)
    M = cv2.getRotationMatrix2D(center, angle, 1.0)
    rotated = cv2.warpAffine(binary_image, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)

    # Convert the processed image to a format suitable for pytesseract
    text_image = Image.fromarray(rotated)

    # Extract text
    extracted_text = pytesseract.image_to_string(text_image)
    print("Extracted Text: ", extracted_text)

    # Save the extracted text to a Word document
    doc = Document()
    doc.add_heading('Extracted Text', level=1)
    doc.add_paragraph(extracted_text)

    # Save the document
    doc.save(output_filename)
    print(f"Text saved to {output_filename}")
